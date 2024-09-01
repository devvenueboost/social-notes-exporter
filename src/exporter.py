import re
from datetime import datetime
from src.auth import get_reddit_instance, get_twitter_api
from src.config import OUTPUT_FOLDER
from praw.models import Comment, Submission
from prawcore.exceptions import Forbidden
import traceback
import tweepy
from tweepy import TweepyException
from docx import Document
import os
from src import db
from src.models.models import RedditPost, RedditComment
import hashlib

def extract_links(text):
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    return re.findall(url_pattern, text)


def get_saved_reddit_notes(reddit):
    saved_notes = []
    try:
        print("Attempting to fetch saved items...")
        user = reddit.user.me()
        print(f"Authenticated as user: {user.name}")

        saved_items = list(user.saved(limit=None))
        print(f"Found {len(saved_items)} saved items on Reddit.")

        if len(saved_items) == 0:
            print("No saved items found. Check if the authenticated user has any saved posts or comments.")
            return None

        for item in saved_items:
            try:
                if isinstance(item, Comment):
                    links = extract_links(item.body)
                    link_text = "\n".join(links) if links else "No links found in comment"
                    saved_notes.append({
                        "type": "Comment",
                        "subreddit": item.subreddit.display_name,
                        "author": item.author.name if item.author else '[deleted]',
                        "permalink": f"https://reddit.com{item.permalink}",
                        "created_utc": item.created_utc,
                        "body": item.body,
                        "links": link_text
                    })
                    print(f"Added comment from r/{item.subreddit.display_name}")
                elif isinstance(item, Submission):
                    links = extract_links(item.selftext) if item.is_self else []
                    link_text = "\n".join(links) if links else "No links found in post body"
                    comments = []
                    item.comments.replace_more(limit=0)  # Flatten comment tree
                    for comment in item.comments.list():
                        comment_links = extract_links(comment.body)
                        comment_link_text = "\n".join(comment_links) if comment_links else "No links found in comment"
                        comments.append({
                            "author": comment.author.name if comment.author else '[deleted]',
                            "body": comment.body,
                            "links": comment_link_text,
                            "created_utc": comment.created_utc,
                            "id": comment.id  # Make sure this line is present
                        })
                    saved_notes.append({
                        "type": "Submission",
                        "subreddit": item.subreddit.display_name,
                        "author": item.author.name if item.author else '[deleted]',
                        "title": item.title,
                        "url": item.url,
                        "created_utc": item.created_utc,
                        "body": item.selftext if item.is_self else '[This is a link post]',
                        "links": link_text,
                        "comments": comments
                    })
                    print(f"Added post from r/{item.subreddit.display_name} with {len(comments)} comments")
                else:
                    print(f"Found unknown item type: {type(item)}")
            except Forbidden:
                print(f"Access denied to item in r/{item.subreddit.display_name}. Skipping.")
            except Exception as e:
                print(f"Error processing Reddit item: {str(e)}")
                print(traceback.format_exc())
                print("Skipping this item and continuing...")

        if not saved_notes:
            print("No accessible saved posts or comments found on Reddit.")
        else:
            print(f"Total Reddit notes collected: {len(saved_notes)}")

    except Exception as e:
        print(f"Error accessing saved Reddit content: {str(e)}")
        print(traceback.format_exc())
        return None

    return saved_notes




def get_twitter_posts(api, username, count=10):
    twitter_notes = []
    try:
        tweets = api.user_timeline(screen_name=username, count=count, tweet_mode="extended")
        for tweet in tweets:
            links = extract_links(tweet.full_text)
            link_text = "\n".join(links) if links else "No links found in tweet"
            twitter_notes.append({
                "author": tweet.user.screen_name,
                "created_at": tweet.created_at.isoformat(),
                "text": tweet.full_text,
                "links": link_text,
                "tweet_url": f"https://twitter.com/{tweet.user.screen_name}/status/{tweet.id}"
            })
        print(f"Collected {len(twitter_notes)} tweets from @{username}")
    except TweepyException as e:
        print(f"Error fetching tweets: {str(e)}")
    return twitter_notes


def save_notes_to_word(notes, platform):
    doc = Document()

    # Add title
    title = f"{platform.capitalize()} Posts Exported - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    doc.add_heading(title, 0)

    for note in notes:
        if platform == "reddit":
            doc.add_heading(f"{'Comment' if note['type'] == 'Comment' else 'Post'} in r/{note['subreddit']}", level=1)
            doc.add_paragraph(f"Author: u/{note['author']}")
            doc.add_paragraph(
                f"Posted on: {datetime.utcfromtimestamp(note['created_utc']).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            if note['type'] == 'Submission':
                doc.add_paragraph(f"Title: {note['title']}")
            doc.add_paragraph(f"URL: {note['permalink' if note['type'] == 'Comment' else 'url']}")
            doc.add_paragraph(f"Content:\n{note['body']}")
            doc.add_paragraph(f"Links found:\n{note['links']}")

            if note['type'] == 'Submission' and 'comments' in note:
                doc.add_heading("Comments", level=2)
                for comment in note['comments']:
                    doc.add_paragraph(f"Comment by: u/{comment['author']}")
                    doc.add_paragraph(
                        f"Posted on: {datetime.utcfromtimestamp(comment['created_utc']).strftime('%Y-%m-%d %H:%M:%S UTC')}")
                    doc.add_paragraph(f"Content:\n{comment['body']}")
                    doc.add_paragraph(f"Links found:\n{comment['links']}")
                    doc.add_paragraph("-" * 30)
        elif platform == "twitter":
            doc.add_heading(f"Tweet by @{note['author']}", level=1)
            doc.add_paragraph(f"Posted on: {note['created_at']}")
            doc.add_paragraph(f"Content:\n{note['text']}")
            doc.add_paragraph(f"Links found:\n{note['links']}")
            doc.add_paragraph(f"Tweet URL: {note['tweet_url']}")

        doc.add_paragraph("=" * 40)

    filename = os.path.join(OUTPUT_FOLDER,
                            f"{platform.capitalize()}_Posts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    print(f"Saved {len(notes)} notes to {filename}")
    return filename


import traceback


def reddit_api():
    print("Starting reddit_api function")
    reddit = get_reddit_instance()
    print("Got Reddit instance")

    if reddit is None:
        print("Reddit instance is None. Aborting.")
        return None

    print("Attempting to get saved Reddit notes")
    reddit_notes = get_saved_reddit_notes(reddit)
    print(f"Got {len(reddit_notes) if reddit_notes else 0} saved notes")

    if reddit_notes:
        print("Attempting to save notes to Word document")
        filename = save_notes_to_word(reddit_notes, "reddit")
        print(f"Saved notes to file: {filename}")

        # Save to database
        try:
            for note in reddit_notes:
                if note['type'] == 'Submission':
                    # Extract post ID from URL
                    post_id = note['url'].split('/')[-3]
                    existing_post = RedditPost.query.get(post_id)
                    if not existing_post:
                        post = RedditPost(
                            id=post_id,
                            title=note['title'],
                            url=note['url'],
                            selftext=note['body'],
                            subreddit=note['subreddit'],
                            created_utc=datetime.utcfromtimestamp(note['created_utc'])
                        )
                        db.session.add(post)
                        db.session.flush()
                        print(f"Added post {post_id} to database")

                    if 'comments' in note:
                        print(f"Found {len(note['comments'])} comments for post {post_id}")
                        for comment in note['comments']:
                            # Generate a unique ID for the comment
                            comment_id = hashlib.md5(f"{post_id}_{comment['author']}_{comment['body']}".encode()).hexdigest()
                            existing_comment = RedditComment.query.get(comment_id)
                            if not existing_comment:
                                db_comment = RedditComment(
                                    id=comment_id,
                                    post_id=post_id,
                                    body=comment['body'],
                                    author=comment['author'],
                                    created_utc=datetime.utcfromtimestamp(comment['created_utc'])
                                )
                                db.session.add(db_comment)
                                print(f"Added comment {comment_id} to post {post_id}")
                            else:
                                print(f"Comment {comment_id} already exists in database")
                    else:
                        print(f"No comments found for post {post_id}")

            db.session.commit()
            print("Saved posts and comments to database")
        except Exception as e:
            print(f"Error saving to database: {str(e)}")
            print(traceback.format_exc())
            db.session.rollback()

        return filename
    else:
        print("No notes found or error occurred while getting notes")
        return None

def twitter_api(username):
    api = get_twitter_api()
    twitter_notes = get_twitter_posts(api, username)
    if twitter_notes:
        filename = save_notes_to_word(twitter_notes, "twitter")
        return filename
    return None